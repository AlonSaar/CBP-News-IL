"""
Generates a Word document (.docx) from the articles store.

Output: output/cbp_hebrew_reviews.docx
Each run regenerates the full file from state/articles.json.

Requires: python-docx  (pip install python-docx)
"""

from __future__ import annotations

import io
import logging
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import requests

import config

logger = logging.getLogger(__name__)

OUTPUT_DIR = config.ROOT / "output"
DOCX_PATH = OUTPUT_DIR / "cbp_hebrew_reviews.docx"


def _download_image(url: str, timeout: int = 10) -> Optional[bytes]:
    """Download image bytes; return None on failure."""
    try:
        resp = requests.get(
            url,
            timeout=timeout,
            headers={"User-Agent": config.CBP_USER_AGENT},
        )
        resp.raise_for_status()
        return resp.content
    except Exception as exc:
        logger.warning("Could not download image %s: %s", url, exc)
        return None


def generate(articles: list[dict]) -> Path:
    """
    Generate the Word document from a list of article dicts.
    Returns the path to the written file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree
    except ImportError:
        raise ImportError(
            "python-docx is required. Run: pip install python-docx"
        )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- Helper: set RTL on a paragraph ---
    def set_rtl(para):
        pPr = para._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        pPr.append(jc)

    # --- Cover page ---
    cover_title = doc.add_paragraph()
    set_rtl(cover_title)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("CBP Hebrew Reviews")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    set_rtl(subtitle)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("סיכומי עיתונות CBP בעברית")
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    set_rtl(meta)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
    run3 = meta.add_run(f"עודכן: {now_str}  |  סה\"כ כתבות: {len(articles)}")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # --- Group by quarter ---
    quarters: dict[str, list[dict]] = {}
    for art in articles:
        q = art.get("quarter", "Q?-????")
        quarters.setdefault(q, []).append(art)

    sorted_quarters = sorted(quarters.keys(), reverse=True)

    for q_idx, q in enumerate(sorted_quarters):
        # Quarter heading
        q_heading = doc.add_paragraph()
        set_rtl(q_heading)
        q_run = q_heading.add_run(q)
        q_run.font.size = Pt(18)
        q_run.font.bold = True
        q_run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)
        # Underline via border
        pPr = q_heading._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "4")
        bot.set(qn("w:color"), "e94560")
        pb.append(bot)
        pPr.append(pb)

        doc.add_paragraph()

        q_articles = sorted(
            quarters[q], key=lambda a: a.get("article_date", ""), reverse=True
        )

        for art_idx, art in enumerate(q_articles):
            # --- Article image ---
            image_url = art.get("image_url", "")
            if image_url:
                img_bytes = _download_image(image_url)
                if img_bytes:
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = img_para.add_run()
                        run_img.add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
                    except Exception as exc:
                        logger.warning("Could not insert image: %s", exc)

            # --- Title ---
            title_para = doc.add_paragraph()
            set_rtl(title_para)
            title_run = title_para.add_run(art.get("title", ""))
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

            # --- Date + Location ---
            meta_para = doc.add_paragraph()
            set_rtl(meta_para)
            date_run = meta_para.add_run(art.get("article_date", "") + "  |  ")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            loc_run = meta_para.add_run(art.get("location", ""))
            loc_run.font.size = Pt(9)
            loc_run.font.bold = True
            loc_run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)

            # --- Crossing type ---
            ct = art.get("crossing_type", "").replace("סוג מעבר:", "").strip()
            if ct:
                ct_para = doc.add_paragraph()
                set_rtl(ct_para)
                ct_run = ct_para.add_run("סוג מעבר: " + ct)
                ct_run.font.size = Pt(9)
                ct_run.font.italic = True
                ct_run.font.color.rgb = RGBColor(0x4c, 0x5f, 0xd5)

            # --- Body ---
            body_para = doc.add_paragraph()
            set_rtl(body_para)
            body_run = body_para.add_run(art.get("body", ""))
            body_run.font.size = Pt(11)
            body_para.paragraph_format.space_after = Pt(6)

            # --- Source URL ---
            src_para = doc.add_paragraph()
            src_run = src_para.add_run(art.get("url", ""))
            src_run.font.size = Pt(8)
            src_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

            # --- Separator ---
            if art_idx < len(q_articles) - 1:
                sep = doc.add_paragraph()
                pPr2 = sep._p.get_or_add_pPr()
                pb2 = OxmlElement("w:pBdr")
                bot2 = OxmlElement("w:bottom")
                bot2.set(qn("w:val"), "single")
                bot2.set(qn("w:sz"), "4")
                bot2.set(qn("w:space"), "4")
                bot2.set(qn("w:color"), "e0e0e0")
                pb2.append(bot2)
                pPr2.append(pb2)
                doc.add_paragraph()

        # Page break between quarters
        if q_idx < len(sorted_quarters) - 1:
            doc.add_page_break()

    doc.save(DOCX_PATH)
    logger.info("Word document written: %s (%d articles)", DOCX_PATH, len(articles))
    return DOCX_PATH
